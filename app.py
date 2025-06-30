# app.py
import chainlit as cl
from dotenv import load_dotenv
import os
import tempfile
from pathlib import Path

# For text extraction
import PyPDF2
from docx import Document as DocxDocument

# For RAG implementation
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.schema import Document as LangchainDocument

try:
    import fitz  # PyMuPDF
    from PIL import Image
    import pytesseract
    import io
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    print("OCR libraries not available. Install with: pip install PyMuPDF pillow pytesseract")

# LLM
from langchain_google_genai import ChatGoogleGenerativeAI

# Load .env file
load_dotenv()

# Setup Gemini LLM
llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.4)

# Global variables to store extracted text and vector store
extracted_text = ""
vectorstore = None
embeddings_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")

def create_chunks_and_vectorstore(text: str, file_name: str):
    """Create text chunks and build vector store for RAG"""
    global vectorstore
    
    print("Creating text chunks for RAG...")
    
    # Initialize text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,  # Smaller chunks for better retrieval
        chunk_overlap=50,  # Some overlap to maintain context
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    # Split text into chunks
    chunks = text_splitter.split_text(text)
    print(f"Created {len(chunks)} text chunks")
    
    # Create LangChain documents from chunks
    documents = []
    for i, chunk in enumerate(chunks):
        doc = LangchainDocument(
            page_content=chunk,
            metadata={
                "source": file_name,
                "chunk_id": i,
                "total_chunks": len(chunks)
            }
        )
        documents.append(doc)
    
    # Create vector store
    vectorstore = FAISS.from_documents(documents, embeddings_model)
    print("Vector store created successfully!")
    
    return vectorstore, len(chunks)

@cl.on_chat_start
async def start():
    """Initialize the chat session"""
    global extracted_text, vectorstore
    extracted_text = ""
    vectorstore = None
    
    welcome_msg = """� **Resume Reviewer Assistant**

I'm your professional resume reviewer! I can analyze your resume and provide detailed feedback to help you land your dream job.

**What I can do:**
📋 **Resume Analysis** - Overall structure and content review
🎯 **Skills Assessment** - Evaluate your skills section
💼 **Experience Review** - Analyze work experience descriptions
🎓 **Education Check** - Review educational background
📈 **ATS Optimization** - Check Applicant Tracking System compatibility
✨ **Improvement Suggestions** - Specific recommendations for enhancement

**Supported formats:**
• PDF files (.pdf) - with OCR support for scanned documents
• Word documents (.docx, .doc)  
• Text files (.txt)

**How to get started:**
1. Upload your resume using drag & drop or the upload feature
2. I'll extract and analyze the content automatically
3. Ask for specific feedback or use the action buttons for quick reviews!

**Upload your resume to begin the review!**
"""
    
    # Add action buttons for common resume review tasks
    actions = [
        cl.Action(name="full_review", value="full_review", description="📋 Complete Resume Review", payload={}),
        cl.Action(name="skills_analysis", value="skills_analysis", description="🎯 Skills Analysis", payload={}),
        cl.Action(name="experience_review", value="experience_review", description="💼 Experience Review", payload={}),
        cl.Action(name="ats_check", value="ats_check", description="📈 ATS Optimization", payload={}),
        cl.Action(name="improvements", value="improvements", description="✨ Improvement Tips", payload={}),
    ]
    
    await cl.Message(content=welcome_msg, actions=actions).send()

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using multiple methods"""
    text_content = ""
    
    # Method 1: PyPDF2
    try:
        print("Trying PyPDF2...")
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
        
        if text_content.strip():
            print(f"PyPDF2 extracted {len(text_content)} characters")
            return text_content
    except Exception as e:
        print(f"PyPDF2 failed: {e}")
    
    # Method 2: OCR with PyMuPDF + Tesseract (if available)
    if HAS_OCR:
        try:
            print("Trying OCR extraction...")
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # First try normal text extraction
                page_text = page.get_text()
                if page_text.strip():
                    text_content += page_text + "\n"
                else:
                    # If no text, try OCR
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Perform OCR
                    ocr_text = pytesseract.image_to_string(img)
                    text_content += ocr_text + "\n"
            
            doc.close()
            
            if text_content.strip():
                print(f"OCR extracted {len(text_content)} characters")
                return text_content
                
        except Exception as e:
            print(f"OCR extraction failed: {e}")
    
    return text_content

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(file_path)
        text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        print(f"DOCX extraction: {len(text_content)} characters")
        return text_content
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        print(f"TXT extraction: {len(text_content)} characters")
        return text_content
    except Exception as e:
        try:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                text_content = f.read()
            print(f"TXT extraction (latin-1): {len(text_content)} characters")
            return text_content
        except Exception as e2:
            print(f"TXT extraction failed: {e2}")
            return ""

@cl.on_message
async def handle_message(message: cl.Message):
    """Handle user messages and file uploads"""
    global extracted_text
    
    # Check if files were uploaded
    if message.elements:
        await handle_file_upload(message.elements)
    else:
        # Check if user is providing text content directly
        if len(message.content) > 500:  # Assume long text might be document content
            await cl.Message(content="📝 **I notice you've sent a long text. Would you like me to treat this as document content?**\n\nType 'YES' to use this text as the document, or ask a question if you already have a document loaded.").send()
            extracted_text = message.content  # Store the text for potential use
            return
        
        # Handle text queries about the uploaded file
        if not extracted_text:
            await cl.Message(content="❌ **No resume uploaded yet.**\n\nPlease upload your resume first:\n• PDF (.pdf)\n• Word documents (.docx, .doc)\n• Text files (.txt)\n\nOnce uploaded, I can provide professional resume feedback and analysis!").send()
            return
        
        # Handle special commands
        if message.content.upper() == 'YES' and len(extracted_text) > 500:
            # User confirmed they want to use the long text as document content
            char_count = len(extracted_text)
            preview = extracted_text[:300] + "..." if len(extracted_text) > 300 else extracted_text
            
            # Create RAG vector store for manual text input
            vectorstore, chunk_count = create_chunks_and_vectorstore(extracted_text, "manual_input.txt")
            
            success_msg = f"""✅ **Resume content loaded successfully!**

📊 **Content processed:** {char_count} characters
🔍 **RAG chunks created:** {chunk_count}

**Content Preview:**
```
{preview}
```

**🎯 Your resume is ready for AI-powered analysis using RAG!**

**Try asking:**
• "Give me a complete resume review"
• "How can I improve my work experience section?"
• "Is my resume ATS-optimized?"
• "What skills should I highlight more?"
"""
            await cl.Message(content=success_msg).send()
            return
        
        # Send query to LLM with the extracted text
        await handle_query_with_llm(message.content)

async def handle_file_upload(elements):
    """Process uploaded files and extract text"""
    global extracted_text
    
    processing_msg = cl.Message(content="📤 **Processing your file...**\n\nExtracting text from the uploaded document...")
    await processing_msg.send()
    
    try:
        for file in elements:
            file_name = getattr(file, 'name', 'unknown_file')
            print(f"Processing file: {file_name}")
            print(f"File has content attribute: {hasattr(file, 'content')}")
            print(f"File content is None: {getattr(file, 'content', None) is None}")
            
            # Save file temporarily - handle both cases: with and without content
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file_name).suffix) as tmp_file:
                if hasattr(file, 'content') and file.content is not None:
                    # File has content - write it
                    if isinstance(file.content, bytes):
                        tmp_file.write(file.content)
                        print(f"Wrote {len(file.content)} bytes to temporary file")
                    else:
                        content_bytes = file.content.encode('utf-8')
                        tmp_file.write(content_bytes)
                        print(f"Wrote {len(content_bytes)} bytes (encoded) to temporary file")
                else:
                    # File has no content - try to get file path or handle differently
                    print(f"File {file_name} has no content attribute, checking for path...")
                    if hasattr(file, 'path') and file.path:
                        print(f"File has path: {file.path}")
                        # Copy from the file path if available
                        try:
                            with open(file.path, 'rb') as src_file:
                                tmp_file.write(src_file.read())
                            print(f"Copied file from path: {file.path}")
                        except Exception as path_error:
                            print(f"Failed to read from path {file.path}: {path_error}")
                            # Create an empty file and let the extraction handle it
                            pass
                    else:
                        print("File has no path either, creating empty temporary file")
                        # Create empty file - extraction will handle the error
                        pass
                
                tmp_file_path = tmp_file.name
            
            # Check file size
            file_size = os.path.getsize(tmp_file_path)
            print(f"Temporary file size: {file_size} bytes")
            
            if file_size == 0:
                # Try alternative approach - maybe the file is accessible through other means
                print(f"Temporary file is empty, trying alternative approaches...")
                
                # Check if we can access the original file through Chainlit's API
                if hasattr(file, 'url') and file.url:
                    print(f"File has URL: {file.url}")
                    # Note: This might not work in all Chainlit versions
                    
                # For now, provide a helpful error message
                error_details = f"""
**Debug Information:**
• File name: {file_name}
• Has content attribute: {hasattr(file, 'content')}
• Content is None: {getattr(file, 'content', None) is None}
• Has path attribute: {hasattr(file, 'path')}
• Path value: {getattr(file, 'path', 'None')}
• Temporary file size: {file_size} bytes

This appears to be a Chainlit file upload issue. The file was uploaded but no content was received.
"""
                
                raise Exception(f"File {file_name} is empty (0 bytes) after upload. {error_details}")
            
            # Extract text based on file type
            file_extension = Path(file_name).suffix.lower()
            
            if file_extension == '.pdf':
                extracted_text = extract_text_from_pdf(tmp_file_path)
            elif file_extension in ['.docx', '.doc']:
                extracted_text = extract_text_from_docx(tmp_file_path)
            elif file_extension == '.txt':
                extracted_text = extract_text_from_txt(tmp_file_path)
            else:
                # Try as text file
                extracted_text = extract_text_from_txt(tmp_file_path)
            
            # Clean up temporary file
            os.unlink(tmp_file_path)
            
            # Check if extraction was successful
            if not extracted_text or not extracted_text.strip():
                raise Exception(f"Could not extract any text from {file_name}. The file might be image-based, password-protected, or corrupted.")
            
            # Create chunks and vector store for RAG
            print("Building RAG vector store...")
            vectorstore, chunk_count = create_chunks_and_vectorstore(extracted_text, file_name)
            
            # Success message with action buttons
            char_count = len(extracted_text)
            preview = extracted_text[:300] + "..." if len(extracted_text) > 300 else extracted_text
            
            # Create action buttons for post-upload analysis
            post_upload_actions = [
                cl.Action(name="full_review", value="full_review", description="📋 Complete Resume Review", payload={}),
                cl.Action(name="skills_analysis", value="skills_analysis", description="🎯 Skills Analysis", payload={}),
                cl.Action(name="experience_review", value="experience_review", description="💼 Experience Review", payload={}),
                cl.Action(name="ats_check", value="ats_check", description="📈 ATS Optimization", payload={}),
                cl.Action(name="improvements", value="improvements", description="✨ Improvement Tips", payload={}),
            ]
            
            success_msg = f"""✅ **Resume uploaded successfully!**

📁 **File:** {file_name}
📊 **Content extracted:** {char_count} characters
📄 **File size:** {file_size} bytes
🔍 **RAG chunks created:** {chunk_count}

**Content Preview:**
```
{preview}
```

**🎯 Your resume is ready for AI-powered analysis using RAG!**

**Use the action buttons below for quick analysis, or ask specific questions:**

**Popular questions:**
• "Give me a complete resume review"
• "How can I better quantify my achievements?"
• "Is my resume ATS-friendly?"
• "What are the biggest improvements I can make?"
• "How can I improve my skills section?"
• "Rate my resume overall"
"""
            
            await cl.Message(content=success_msg, actions=post_upload_actions).send()
            processing_msg.content = "📤 **File processing complete!** ✅"
            await processing_msg.update()
            
            print(f"Successfully extracted {char_count} characters from {file_name}")
            return
            
    except Exception as e:
        error_msg = f"""❌ **File processing failed**

**Error:** {str(e)}

**Troubleshooting:**
• Check the console/terminal for detailed debug information
• Try uploading a smaller file first to test
• Make sure the file is not password-protected or corrupted
• For PDFs: Try converting to text format first if it's a scanned document

**Alternative solution:**
If file upload continues to fail, you can:
1. Copy the text from your document manually
2. Paste it directly in the chat
3. I'll analyze the text content directly

**Supported formats:**
• PDF (.pdf)
• Word documents (.docx, .doc)
• Text files (.txt)
"""
        processing_msg.content = error_msg
        await processing_msg.update()
        print(f"File processing error: {e}")
        import traceback
        traceback.print_exc()

async def handle_query_with_llm(query: str):
    """Send the extracted text and user query to the LLM for resume analysis using RAG"""
    global extracted_text, vectorstore
    
    thinking_msg = cl.Message(content="🔍 **Analyzing your resume with RAG...**\n\nRetrieving most relevant sections and providing professional feedback...")
    await thinking_msg.send()
    
    try:
        # If no vector store exists, create one from the extracted text
        if not vectorstore and extracted_text:
            print("No vector store found, creating one...")
            vectorstore, _ = create_chunks_and_vectorstore(extracted_text, "uploaded_resume")
        
        if not vectorstore:
            thinking_msg.content = "❌ **Error:** No resume content available for analysis. Please upload a resume first."
            await thinking_msg.update()
            return
        
        # Use RAG to retrieve relevant document chunks
        relevant_docs = vectorstore.similarity_search(query, k=4)  # Get top 4 most relevant chunks
        
        # Combine relevant chunks into context
        relevant_context = "\n\n".join([doc.page_content for doc in relevant_docs])
        
        print(f"Retrieved {len(relevant_docs)} relevant chunks for query: {query}")
        for i, doc in enumerate(relevant_docs):
            print(f"Chunk {i+1} metadata: {doc.metadata}")
        
        # Create resume-focused prompt with RAG context
        prompt = f"""You are a professional resume reviewer and career coach. A job seeker has uploaded their resume and is asking for feedback.

MOST RELEVANT RESUME SECTIONS (Retrieved via RAG):
{relevant_context}

USER QUESTION: {query}

Please provide professional, actionable advice based on the relevant resume sections above. Focus on:
- Specific improvements that can be made
- Industry best practices
- ATS optimization where relevant
- Career advancement opportunities
- Concrete examples and suggestions

If the question is about a specific section (skills, experience, education), provide detailed analysis of that area. Always be constructive and helpful in your feedback.

Note: Your analysis is based on the most relevant sections of the resume that match the user's query."""

        # Get response from LLM
        response = llm.invoke(prompt)
        
        # Update message with response
        thinking_msg.content = f"💼 **Professional Feedback (RAG-Enhanced):**\n\n{response.content}\n\n---\n*📊 Analysis based on {len(relevant_docs)} most relevant resume sections*"
        await thinking_msg.update()
        
    except Exception as e:
        thinking_msg.content = f"❌ **Error:** Sorry, I encountered an issue: {str(e)}"
        await thinking_msg.update()
        print(f"LLM query error: {e}")
        import traceback
        traceback.print_exc()

# Action callbacks for resume review buttons
@cl.action_callback("full_review")
async def on_full_review(action):
    if not extracted_text:
        await cl.Message(content="❌ Please upload your resume first!").send()
        return
    await handle_resume_analysis("full_review")

@cl.action_callback("skills_analysis")
async def on_skills_analysis(action):
    if not extracted_text:
        await cl.Message(content="❌ Please upload your resume first!").send()
        return
    await handle_resume_analysis("skills_analysis")

@cl.action_callback("experience_review")
async def on_experience_review(action):
    if not extracted_text:
        await cl.Message(content="❌ Please upload your resume first!").send()
        return
    await handle_resume_analysis("experience_review")

@cl.action_callback("ats_check")
async def on_ats_check(action):
    if not extracted_text:
        await cl.Message(content="❌ Please upload your resume first!").send()
        return
    await handle_resume_analysis("ats_check")

@cl.action_callback("improvements")
async def on_improvements(action):
    if not extracted_text:
        await cl.Message(content="❌ Please upload your resume first!").send()
        return
    await handle_resume_analysis("improvements")

async def handle_resume_analysis(analysis_type: str):
    """Handle different types of resume analysis using RAG"""
    global extracted_text, vectorstore
    
    thinking_msg = cl.Message(content="🔍 **Analyzing your resume with RAG...**\n\nRetrieving relevant sections and generating professional feedback...")
    await thinking_msg.send()
    
    try:
        # If no vector store exists, create one from the extracted text
        if not vectorstore and extracted_text:
            print("No vector store found, creating one...")
            vectorstore, _ = create_chunks_and_vectorstore(extracted_text, "uploaded_resume")
        
        if not vectorstore:
            thinking_msg.content = "❌ **Error:** No resume content available for analysis. Please upload a resume first."
            await thinking_msg.update()
            return
        
        # Define search queries for different analysis types
        search_queries = {
            "full_review": "resume structure format contact information summary skills work experience education",
            "skills_analysis": "skills technical skills programming languages software tools abilities competencies",
            "experience_review": "work experience job responsibilities achievements accomplishments employment history",
            "ats_check": "keywords job titles skills industry terms technical terms",
            "improvements": "resume content improvements suggestions formatting structure experience skills"
        }
        
        # Get relevant chunks for the specific analysis type
        search_query = search_queries.get(analysis_type, "resume analysis feedback")
        relevant_docs = vectorstore.similarity_search(search_query, k=6)  # Get more chunks for comprehensive analysis
        
        # Combine relevant chunks into context
        relevant_context = "\n\n".join([doc.page_content for doc in relevant_docs])
        
        print(f"Retrieved {len(relevant_docs)} relevant chunks for {analysis_type} analysis")
        
        # Different prompts for different analysis types with RAG context
        prompts = {
            "full_review": f"""You are a professional resume reviewer and career coach. Please provide a comprehensive review based on the most relevant sections:

MOST RELEVANT RESUME SECTIONS (Retrieved via RAG):
{relevant_context}

Please provide a detailed analysis covering:
1. **Overall Structure & Format** - Is it well-organized and professional?
2. **Contact Information** - Is it complete and professional?
3. **Professional Summary/Objective** - Is it compelling and targeted?
4. **Skills Section** - Are skills relevant and well-presented?
5. **Work Experience** - Are achievements quantified and impactful?
6. **Education** - Is it appropriately presented?
7. **Overall Strengths** - What works well?
8. **Areas for Improvement** - Specific suggestions for enhancement
9. **Overall Rating** - Rate from 1-10 with explanation

Format your response with clear headings and actionable feedback based on the relevant sections provided.""",

            "skills_analysis": f"""You are a skills assessment expert. Please analyze the skills-related content from these relevant sections:

MOST RELEVANT RESUME SECTIONS (Retrieved via RAG):
{relevant_context}

Focus on:
1. **Technical Skills** - Are they current and relevant?
2. **Soft Skills** - Are they well-integrated and demonstrated?
3. **Skills Organization** - How are skills categorized and presented?
4. **Market Relevance** - Are these skills in demand?
5. **Missing Skills** - What important skills might be missing?
6. **Skills Presentation** - How can skills be better showcased?
7. **Recommendations** - Specific skills to add or emphasize

Provide actionable advice for improving the skills section based on the relevant content.""",

            "experience_review": f"""You are a career counselor specializing in work experience optimization. Please review the experience-related content:

MOST RELEVANT RESUME SECTIONS (Retrieved via RAG):
{relevant_context}

Analyze:
1. **Job Descriptions** - Are they achievement-focused or task-focused?
2. **Quantification** - Are accomplishments backed by numbers/metrics?
3. **Action Verbs** - Are strong action verbs used effectively?
4. **Relevance** - Is experience relevant to career goals?
5. **Progression** - Does it show career growth?
6. **Impact Statements** - Are contributions to employers clear?
7. **Gaps or Issues** - Any concerning patterns?
8. **Improvement Suggestions** - How to strengthen each role description

Provide specific examples of how to rewrite weak bullet points based on the content provided.""",

            "ats_check": f"""You are an ATS (Applicant Tracking System) optimization expert. Please review the relevant sections for ATS compatibility:

MOST RELEVANT RESUME SECTIONS (Retrieved via RAG):
{relevant_context}

Check for:
1. **Keyword Optimization** - Are industry keywords present?
2. **Format Compatibility** - Is the format ATS-friendly?
3. **Section Headers** - Are they standard and recognizable?
4. **Contact Information** - Is it properly formatted?
5. **Skills Keywords** - Are important skills mentioned?
6. **Job Title Alignment** - Do titles match industry standards?
7. **Acronyms & Abbreviations** - Are both forms included?
8. **ATS Score** - Estimated ATS compatibility (1-10)

Provide specific recommendations to improve ATS compatibility and keyword optimization.""",

            "improvements": f"""You are a resume improvement specialist. Please provide specific, actionable improvement suggestions based on the relevant sections:

MOST RELEVANT RESUME SECTIONS (Retrieved via RAG):
{relevant_context}

Provide:
1. **Quick Wins** - Easy changes that make big impact
2. **Content Improvements** - How to strengthen content
3. **Formatting Suggestions** - Visual and structural improvements
4. **Language Enhancement** - Better word choices and phrasing
5. **Missing Elements** - What important sections or info is missing?
6. **Industry Alignment** - How to better match industry expectations
7. **Priority Action Items** - Top 5 changes to make first
8. **Before/After Examples** - Show how to rewrite specific sections

Focus on practical, implementable advice based on the content provided."""
        }
        
        prompt = prompts.get(analysis_type, prompts["full_review"])
        
        # Get response from LLM
        response = llm.invoke(prompt)
        
        # Create title based on analysis type
        titles = {
            "full_review": "📋 **Complete Resume Review (RAG-Enhanced)**",
            "skills_analysis": "🎯 **Skills Analysis (RAG-Enhanced)**",
            "experience_review": "💼 **Experience Review (RAG-Enhanced)**",
            "ats_check": "📈 **ATS Optimization Report (RAG-Enhanced)**",
            "improvements": "✨ **Improvement Recommendations (RAG-Enhanced)**"
        }
        
        title = titles.get(analysis_type, "📋 **Resume Analysis (RAG-Enhanced)**")
        
        # Update message with response
        thinking_msg.content = f"{title}\n\n{response.content}\n\n---\n*📊 Analysis based on {len(relevant_docs)} most relevant resume sections*"
        await thinking_msg.update()
        
    except Exception as e:
        thinking_msg.content = f"❌ **Error:** Sorry, I encountered an issue during analysis: {str(e)}"
        await thinking_msg.update()
        print(f"Resume analysis error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    print("Starting Simple File Chat Assistant...")
    print(f"OCR capabilities: {'Available' if HAS_OCR else 'Not available'}")
    if not HAS_OCR:
        print("To enable OCR for scanned PDFs, install: pip install PyMuPDF pillow pytesseract")
